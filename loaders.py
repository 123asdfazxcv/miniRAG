import csv
import json
import logging
import re
from abc import ABC,abstractmethod
from pathlib import Path
from typing import Optional

from minirag import Document

logger = logging.getLogger("RAG.Loader")
def _read_text(path:str) -> str:
    encodings = ["utf-8","gbk","gb2312","gb18030","latin-1"]
    for enc in encodings:
        try:
            with open(path,encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path,encoding="latin-1") as f:
        return f.read()
class BaseLoader(ABC):

    def load(self,file_path:str,filename:Optional[str]=None)->list[Document]:
        ...

#########################
# Markdown Loader
#########################
class MarkdownLoader(BaseLoader):
    HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$',re.MULTILINE)

    def load(self,file_path:str,filename:Optional[str]=None) ->list[Document]:

        text = _read_text(file_path)
        source = filename or Path(file_path).name
        matches = list(self.HEADING_PATTERN.finditer(text))

        if not matches:
            return [Document(content=text.strip(),metadata={"source":source})]
        docs:list[Document] = []
        for i,match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1< len(matches) else len(text)
            section_text = text[start:end].strip()
            if section_text:
                docs.append(Document(
                    content=section_text,
                    metadata={
                        "source":source,
                        "heading":match.group(2).strip(),
                        "level":len(match.group(1)),
                    }
                ))
        return docs
class PDFLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None)->list[Document]:
        import pdfplumber
        source = filename or Path(file_path).name
        docs:list[Document] = []
        with pdfplumber.open(file_path) as pdf:
            total = len(pdf.pages)
            for i ,page in enumerate(pdf.pages,1):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        content=text.strip(),
                        metadata={"source":source,"page":i,"total_pages":total}
                    ))
        return docs

class CSVLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None) ->list[Document]:
        source = filename or Path(file_path).name
        docs:list[Document] = []
        text = _read_text(file_path)
        import io
        reader = csv.DictReader(io.StringIO(text))
        for row_num,row in enumerate(reader,1):
            parts = [f"{k}:{v}" for k,v in row.items() if v is not None and v != ""]
            if parts:
                docs.append(Document(
                    content=",".join(parts),
                    metadata={"source":source,"row":row_num,"headers":list(row.keys())}
                ))
        return docs

class ExcelLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None) ->list[Document]:
        import openpyxl
        source = filename or Path(file_path).name
        docs:list[Document] = []
        wb = openpyxl.load_workbook(file_path,read_only=True,data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers:list[str] = []
            for row_idx ,row in enumerate(ws.iter_rows(values_only=True),1):
                cells = [str(c) if c is not None else "" for c in row]
                if row_idx == 1:
                    headers = cells
                    continue
                parts = [f"{h}:{v}" for h ,v in zip(headers,cells) if v]
                if parts:
                    docs.append(Document(
                        content=",".join(parts),
                        metadata={"source":source,"sheet":sheet_name,"row":row_idx}
                    ))
        wb.close()
        return docs
class JSONLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None) ->list[Document]:
        source = filename or Path(file_path).name
        text = _read_text(file_path)
        data = json.loads(text)
        if isinstance(data,list):
            return [
                Document(content=str(item),metadata={"source":source,"index":i})
                for i,item in enumerate(data) if item
            ]
        elif isinstance(data,dict):
            return [
                Document(content=str(v),metadata={"source":source,"index":k})
                for k,v in data.items() if v
            ]
        return [Document(content=str(data),metadata={"source":source})]
class DocxLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None) ->list[Document]:
        from docx import Document as DocxDocument
        source = filename or Path(file_path).name
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return [Document(content=content,metadata={"source":source})]

class TextLoader(BaseLoader):
    def load(self,file_path:str,filename:Optional[str]=None)->list[Document]:
        text = _read_text(file_path)
        source = filename or Path(file_path).name
        return [Document(content=text.strip(),metadata={"source":source})]

class DocumentLoaderFactory:
    EXTENSION_MAP:dict[str,type[BaseLoader]] = {
        ".md":MarkdownLoader,".markdown":MarkdownLoader,
        ".pdf":PDFLoader,
        ".csv":CSVLoader,
        ".xlsx":ExcelLoader,".xls":ExcelLoader,
        ".json":JSONLoader,
        ".txt":TextLoader,".py":TextLoader,".js":TextLoader,
        ".ts":TextLoader,".yaml":TextLoader,".yml":TextLoader,
        ".xml":TextLoader,".html":TextLoader,".htm":TextLoader,
        ".ini":TextLoader,".docx":DocxLoader
    }

    @classmethod
    def get_loader(cls,file_path:str) -> BaseLoader:
        ext =Path(file_path).suffix.lower()
        loader_cls = cls.EXTENSION_MAP.get(ext,TextLoader)
        if loader_cls is None:
            logger.warning(f"未识别扩展名'{ext}',使用TextLoader兜底")
            loader_cls = TextLoader
        return loader_cls()